import pdfplumber
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document
import os
import html
from pathlib import Path

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

class PDFExtractor:
    def __init__(self, pdf_path):
        self.pdf_path = pdf_path

    def _extract_with_plumber(self, page):
        try:
            text = page.extract_text(x_tolerance=1, y_tolerance=1, layout=True)
            return text if text else ""
        except Exception as e:
            print(f"Error extracting with pdfplumber on page {page.page_number}: {e}")
            return ""

    def _extract_with_ocr(self, page):
        try:
            pil_image = page.to_image(resolution=300).original
            text = pytesseract.image_to_string(pil_image, lang='eng')
            return text if text else ""
        except pytesseract.TesseractNotFoundError:
            print("Tesseract is not installed or not in your PATH. Please install it.")
            return ""
        except Exception as e:
            print(f"Error extracting with OCR on page {page.page_number}: {e}")
            return ""
    
    def extract_data_layout_preserved(self):
        extracted_data = {}
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                for page in pdf.pages:
                    plumber_text = self._extract_with_plumber(page)
                    
                    if plumber_text:
                        extracted_data[page.page_number] = plumber_text
                    else:
                        ocr_text = self._extract_with_ocr(page)
                        extracted_data[page.page_number] = ocr_text

            print("PDF extraction completed successfully!")
            return extracted_data
        except FileNotFoundError:
            print(f"Error: The file at {self.pdf_path} was not found.")
            return {}
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            return {}

    def extract_elements_for_html(self):
        extracted_elements = {}
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                for page in pdf.pages:
                    elements = []
                    words = page.extract_words(x_tolerance=1, y_tolerance=1)
                    if words:
                        for word in words:
                            elements.append({
                                'text': word['text'],
                                'x0': word['x0'], 'y0': word['top'],
                                'x1': word['x1'], 'y1': word['bottom'],
                                'source': 'plumber'
                            })
                    else:
                        pil_image = page.to_image(resolution=300).original
                        data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT)
                        for i in range(len(data['text'])):
                            text = str(data['text'][i]).strip()
                            conf = float(data['conf'][i]) if data['conf'][i] != '-1' else 0
                            if text and conf > 50:
                                x, y, w, h = int(data['left'][i]), int(data['top'][i]), int(data['width'][i]), int(data['height'][i])
                                elements.append({
                                    'text': text,
                                    'x0': x, 'y0': y,
                                    'x1': x + w, 'y1': y + h,
                                    'source': 'ocr'
                                })
                    extracted_elements[page.page_number] = elements
            return extracted_elements
        except FileNotFoundError:
            print(f"Error: The file at {self.pdf_path} was not found.")
            return {}
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            return {}

    def save_to_docx(self, data, output_filename):
        document = Document()
        for page_num, content in data.items():
            document.add_heading(f"Page {page_num}", level=1)
            document.add_paragraph(content)
            document.add_page_break()
        document.save(output_filename)
        print(f"Data saved to {output_filename} successfully!")

    def save_to_txt(self, data, output_filename):
        with open(output_filename, 'w', encoding='utf-8') as f:
            for page_num, content in data.items():
                f.write(f"--- Page {page_num} ---\n")
                f.write(content)
                f.write("\n\n")
        print(f"Data saved to {output_filename} successfully!")

    def save_to_html(self, data, output_filename):
        html_content = ["<html><head><meta charset='utf-8'><title>PDF Extraction</title></head><body>"]
        html_content.append("<style>")
        html_content.append("body { font-family: Arial, sans-serif; margin: 0; padding: 20px; background-color: #f0f0f0; }")
        html_content.append(".page { position: relative; margin: 20px auto; background-color: #fff; border: 1px solid #ccc; box-shadow: 0 0 10px rgba(0,0,0,0.1); padding: 50px; box-sizing: border-box; }")
        html_content.append(".text-element { position: absolute; font-size: 12px; white-space: pre-wrap; margin: 0; padding: 0; }")
        html_content.append("</style>")

        for page_num, elements in data.items():
            if not elements:
                continue

            max_x = max(e['x1'] for e in elements)
            max_y = max(e['y1'] for e in elements)

            html_content.append(f'<div class="page" style="width: {max_x}px; height: {max_y}px;">')
            
            for element in elements:
                safe_text = html.escape(element['text'])
                style = f"left: {element['x0']}px; top: {element['y0']}px;"
                style += f"width: {element['x1'] - element['x0']}px;"
                style += f"height: {element['y1'] - element['y0']}px;"
                
                html_content.append(f'<p class="text-element" style="{style}">{safe_text}</p>')
            
            html_content.append('</div>')

        html_content.append("</body></html>")
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
            
        print(f"Data saved to {output_filename} successfully!")

def main():
    pdf_path = input("Please enter the path to the PDF file: ").strip()
    
    if not os.path.exists(pdf_path):
        print(f"Error: The file '{pdf_path}' does not exist.")
        return

    extractor = PDFExtractor(pdf_path)
    layout_data = extractor.extract_data_layout_preserved()
    html_elements_data = extractor.extract_elements_for_html()

    if not layout_data:
        print("No data was extracted. Exiting.")
        return

    output_filename = input("Enter the desired output filename (e.g., 'report'): ").strip()
    output_format = input("Choose output format (docx, txt, html): ").strip().lower()

    if output_format == 'docx':
        extractor.save_to_docx(layout_data, f"{output_filename}.docx")
    elif output_format == 'txt':
        extractor.save_to_txt(layout_data, f"{output_filename}.txt")
    elif output_format == 'html':
        extractor.save_to_html(html_elements_data, f"{output_filename}.html")
    else:
        print("Invalid output format chosen. Supported formats are: docx, txt, html.")

if __name__ == "__main__":
    main()
