import pdfplumber
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document
from fpdf import FPDF
import os

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

class PDFExtractor:
    """
    A class to extract text from PDF documents while preserving layout.
    It uses a hybrid approach, combining pdfplumber for digital PDFs and Tesseract for scanned/image-based PDFs.
    """

    def __init__(self, pdf_path):
        """
        Initializes the PDFExtractor with the path to the PDF file.
        :param pdf_path: Path to the PDF file.
        """
        self.pdf_path = pdf_path

    def _extract_with_plumber(self, page):
        """
        Extracts text from a single page using pdfplumber, preserving layout.
        :param page: A pdfplumber page object.
        :return: Extracted text as a string.
        """
        try:
            text = page.extract_text(x_tolerance=1, y_tolerance=1, layout=True)
            return text if text else ""
        except Exception as e:
            print(f"Error extracting with pdfplumber on page {page.page_number}: {e}")
            return ""

    def _extract_with_ocr(self, page):
        """
        Performs OCR on a page that is likely an image-based PDF.
        :param page: A pdfplumber page object.
        :return: Extracted text as a string.
        """
        try:
            pil_image = page.to_image(resolution=300).original
            text = pytesseract.image_to_string(pil_image, lang='eng')
            return text if text else ""
        except pytesseract.TesseractNotFoundError:
            print("Tesseract is not installed or not in your PATH. Please install it.")
            return ""
        except Exception as e:
            print(f"Error extracting with OCR on page {page.page_number}: {e}")
            return ""

    def extract_data(self):
        """
        Main method to extract data from the entire PDF.
        :return: A dictionary with page numbers as keys and extracted text as values.
        """
        extracted_data = {}
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                for page in pdf.pages:
                    plumber_text = self._extract_with_plumber(page)
                    
                    if plumber_text:
                        extracted_data[page.page_number] = plumber_text
                    else:
                        ocr_text = self._extract_with_ocr(page)
                        extracted_data[page.page_number] = ocr_text

            print("PDF extraction completed successfully!")
            return extracted_data
        except FileNotFoundError:
            print(f"Error: The file at {self.pdf_path} was not found.")
            return {}
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            return {}
    
    def save_to_docx(self, data, output_filename):
        """Saves the extracted data to a DOCX file."""
        document = Document()
        for page_num, content in data.items():
            document.add_heading(f"Page {page_num}", level=1)
            document.add_paragraph(content)
            document.add_page_break()
        document.save(output_filename)
        print(f"Data saved to {output_filename} successfully! ")

    def save_to_txt(self, data, output_filename):
        """Saves the extracted data to a TXT file."""
        with open(output_filename, 'w', encoding='utf-8') as f:
            for page_num, content in data.items():
                f.write(f"--- Page {page_num} ---\n")
                f.write(content)
                f.write("\n\n")
        print(f"Data saved to {output_filename} successfully! ")

    def save_to_html(self, data, output_filename):
        """Saves the extracted data to an HTML file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for page_num, content in data.items():
            pdf.write(5, f"<h1>Page {page_num}</h1>")
            # Encode the text to handle special characters correctly
            pdf.multi_cell(0, 5, content.encode('latin-1', 'replace').decode('latin-1'))
            pdf.add_page()
        pdf.output(output_filename, 'F')
        print(f"Data saved to {output_filename} successfully! ")

def main():
    """
    Main function to handle user interaction, extraction, and file saving.
    """
    pdf_path = input("Please enter the path to the PDF file: ").strip()
    
    if not os.path.exists(pdf_path):
        print(f"Error: The file '{pdf_path}' does not exist.")
        return

    extractor = PDFExtractor(pdf_path)
    data = extractor.extract_data()

    if not data:
        print("No data was extracted. Exiting.")
        return

    output_filename = input("Enter the desired output filename (e.g., 'report'): ").strip()
    output_format = input("Choose output format (docx, txt, html): ").strip().lower()

    if output_format == 'docx':
        extractor.save_to_docx(data, f"{output_filename}.docx")
    elif output_format == 'txt':
        extractor.save_to_txt(data, f"{output_filename}.txt")
    elif output_format == 'html':
        extractor.save_to_html(data, f"{output_filename}.html")
    else:
        print("Invalid output format chosen. Supported formats are: docx, txt, html.")

if __name__ == "__main__":
    main()